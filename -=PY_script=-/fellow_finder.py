import re
from itertools import groupby
import configparser
import os
import pathlib
from datetime import datetime
import json
import collections
import glob

import requests
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Cm
import networkx as nx
import matplotlib.pyplot as plt
import numpy as np


class FellowFinder:
    def __init__(self):
        # входная директория
        self.in_dir = 'c:/in'
        # выходная директория(результирующая)
        self.out_dir = ''
        # директория для хранения инф-аналит. документ :))
        self.iad_dir = ''
        # кол-во общих поездок после человек считается попутчиков
        # анализируемой персоны
        self.trips_number = 3
        # кол-во дней пребывания за границей, по которому сообщаем, что
        # человек выезжал.
        self.days_aboard = 0
        # соберем шаблон для поиска заранее - так быстрее работать будет, шаблон для поиска инфы о перелетах
        self.re_person = re.compile(
            '(.{5,})\s(\d{2}\.\d{2}\.\d{4})\s(\d{2}\x3A\d{2})\s\x2d\s(.{3,})\n.{0,4}\w{4,5}\s№([\d\w]{3,})[\n\s]{,1}')
            # '(.{5,})\s(\d{2}\.\d{2}\.\d{4})\s(\d{2}\x3A\d{2})\s\x2d\s(.{3,})\n\s\s\w{4,5}\s№([\d\w]{3,})[\n\s]{,1}')
        # шаблон для поиска строк попутчиков
        self.re_travelers = re.compile('(.{5,})\((\d{2}\.\d{2}\.\d{4})')

        # имя текущего файла который обрабатываем
        self.current_fn = ''
        # буффер полученный после чтения
        self.buffer = ''
        # имя персоны, которое передается в первой строке
        self.person_fio = ''
        # список поездок наблюдаемой персоны
        # [('Москва (Шереметьево)(SVО)','22.11.2016','17:45','Минск(Национальный аэропорт Минск)(МSQ)','SU1843')..
        self.person_trips = []
        # список направлений которыек посещала наблюдаемая персона
        # получим список о перелетах вида - 'Новосибирск', 'Москва', '23.01.2015'
        # это нужно для работы с графикой и проверки пребывания за границей
        self.trips_cities = []
        # список с заграничными городами
        self.abroad_trips = []
        # словарь в котором будем хранить города, страну которых уже установили,
        #  ключ - город, значение по ключу 1 - рф
        self.checked_cities = {}
        # список поездок пользователя вида [(Москва, Псков, 11)
        # список отсортирован в порядке убывания, повторы удалены
        self.trips_sorted = []
        # список всех поездок персоны
        # src, dst, date, time
        self.trips_all = []
        # словарь с количество посещенний каждого города
        self.cities_dict = {}
        # отсортированный словарь self.cities_dict, но это уже список вида [(Псков, 11),(Орел, 11), ]
        self.cities_sorted = []

        # отсортированный словарь self.cities_dict, но это уже список вида [(Псков, 11),(Орел, 11), ]
        self.cities_sorted_holidays = []
        # отсортированный словарь self.cities_dict, но это уже список вида [(Псков, 11),(Орел, 11), ]
        self.cities_sorted_weekends = []

        # словарь для хранения информации о совместных поездках с попутчиками
        # key - fio попутчика
        # value - список поездок с датой
        # например: fellows_trips['Иванов Иван Иванович'] = [['21.03.12', 'МСК - СПБ'], ...]
        self.fellows_trips = dict()
        # список поездок в праздничные дни 30.12 - 01.01, 22-24.02, 7-9.03, 8-12.05
        # [[src, dst, cnt]...
        self.trips_sorted_holidays = []
        # (src, dst, date, time)
        self.trips_holidays_all = []
        # список поездок в выходные дни пятница, суббота, воскресенье
        # [[src, dst, cnt]...
        self.trips_sorted_weekends = []
        # (src, dst, date, time)
        self.trips_weekends_all = []
        self.days_dict = dict()
        self.days_dict[5] = 'Пятница'
        self.days_dict[6] = 'Суббота'
        self.days_dict[7] = 'Воскресенье'

        # номер дня в который совершается много поездок
        self.day_ind = -1
        # название дня в который совершаются поездки
        self.day = ''
        # отсортированный список поездок в топ день
        # [(Псков, Питер, 12)....]
        self.trips_sorted_topday = []

    def clear(self):
        """затрем все переменные, списки и словари"""
        self.trips_sorted_topday = []
        self.day = ''
        self.day_ind = -1
        self.current_fn = ''
        self.buffer = ''
        self.person_fio = ''
        self.person_trips.clear()
        self.trips_cities.clear()
        self.abroad_trips.clear()
        self.checked_cities = {}
        self.trips_sorted.clear()
        self.trips_all.clear()
        self.cities_dict = {}
        self.cities_sorted.clear()
        self.fellows_trips = dict()
        self.trips_sorted_holidays.clear()
        self.trips_sorted_weekends.clear()
        self.cities_sorted_holidays.clear()
        self.cities_sorted_weekends.clear()
        self.trips_holidays_all.clear()
        self.trips_weekends_all.clear()

    def load_data(self):
        """считаем данные в зависимости от файла"""
        result = True
        try:
            # получим расширение файла
            _, ext = os.path.splitext(self.current_fn)
            if ext == '.docx':
                # читаем из docx
                doc = Document(self.current_fn)
                doc_raws = [p.text for p in doc.paragraphs]
                # считаем имя анализируемой персоны
                self.person_fio = doc_raws[0] if len(doc_raws) > 0 else ''

                # если фио не получено из первой строки
                # то, возможно оно во второй строке
                if self.person_fio == '':
                    if len(doc_raws) > 1:
                        if doc_raws[1].find(' № ') > 0:
                            self.person_fio = doc_raws[1]

                # если все же имя персоны не получено, вставим имя
                if self.person_fio == '':
                    self.person_fio = pathlib.Path(self.current_fn).stem

                self.buffer = '\n'.join(doc_raws)
            else:
                # читаем обычный текст
                with open(self.current_fn, 'r', encoding='utf-8') as ff:
                    # считаем имя анализируемой персоны
                    self.person_fio = ff.readline()
                    self.buffer = ff.read()
        except Exception as e:
            print(datetime.now(), e)
            result = False

        return result

    def get_tripinfo(self, trip_date):
        """Получим информацию о рейсе по дате
        :param trip_date: дата интересуемого рейса
        :return: Москва(Шереметьево)(ШРМ) - Архангельск(Талаги)(АХГ)  Рейс\поезд №5Н118 22.04.2010 20:10 """
        result = ''
        # получим рейсы в этот день
        for src, date, time, dst, reys in filter(lambda x: x[1] == trip_date, self.person_trips):
            # если такого номера рейса еще нет в результатах
            if result.find(f'№{reys}') == -1:
                # добавочный символ
                addsym = ',\n' if len(result) > 0 else ''
                result += f'{addsym}{src} - {dst}{addsym} Рейс/Поезд №{reys} {date} {time}'
        return result

    def find_fellows(self):
        """
        Главноя функция обрабокти загруженных данных
        :return: в результате будут собраны списки self.csv_data и self.word_data с инфой о попутчиках
        """

        # получим список рейсов персоны - [(from, date, time, to, reys)]
        # [('Москва (Шереметьево)(SVО)','22.11.2016','17:45','Минск(Национальный аэропорт Минск)(МSQ)','SU1843')..
        self.person_trips = self.re_person.findall(self.buffer)
        # ифна по попутчиках, в этот список попадет и персона, но на результат это не повлияет : [(ФИО инфа, дата), ]
        # [('  DАNILОVА DАRIААМRS 0708 007588, 07 08 007588, ПС0708 007588 ', '08.04.2012'), .... ]
        travelers_all = self.re_travelers.findall(self.buffer)

        # получим даты в которые персона ездила
        # ['22.11.2016', '22.10.2016' ...
        dates = [dt for src, dt, time, dst, reys in self.person_trips]
        # получим список попутчиков кто был именно в эти даты
        # [ ('  DIDYК МАIIА МRS 643 728109, Р643 728109 ', '24.08.2011'), ..]
        travelers = ((name, dt) for name, dt in travelers_all if dt in dates)
        # сортируем список для дальнейшей группировки
        travelers = sorted(travelers, key=lambda v: v[0])

        # группируем попутчиков по имени и смотрим кто ездил в те же даты более 2ух раз
        # v[0] - это имя попутчика
        # fellow_name - это имя попутчика
        # trips - группа поездок попутчика с именем fellow_name
        # # [('  DIDYК МАIIА МRS 643 728109, Р643 728109 ', '24.08.2011') .....
        for fellow_name, trips in groupby(travelers, key=lambda v: v[0]):
            trips = tuple(trips)
            # если количество поездок в эту дату больше того, что указано в настройках
            if len(trips) >= self.trips_number:
                # информация о поездках попутчика
                # trips_info = ''
                for i, (name, trip_date) in enumerate(trips):
                    # ГАБОВ АЛЕКСЕЙ ВАЛЕРИЕВИЧ 46 01 916867 ; 22.04.2010; информация о рейсе
                    #                                       ; 20.05.2010; информация о рейсе
                    #                                       ; 23.06.2010; информация о рейсе
                    # если это первый элемент про попутчика, добавим его в словарь
                    if i == 0:
                        self.fellows_trips[name] = []
                    # добавим инфу о поездке с конкретным попутчиком
                    self.fellows_trips[name].append((trip_date, self.get_tripinfo(trip_date)))

    @staticmethod
    def replace_bad_words(city, for_graph=False):
        """Удалим слова из городов, которые яндекс не воспринимает,
        еще перенос - на след строку"""
        bad_words = ('-ГЛАВН.', '-ПАСС', '-ЛАДОЖ', '-ГЛАВН')
        for badword in bad_words:
            if badword in city:
                city = city.replace(badword, '')

        if for_graph and '-' in city:
            city = city.replace('-', '-\n')
        return city

    def get_cities(self, for_graph=False):
        """Переделаем список  из self.person_trips:
        'Новосибирск (Толмачево)(ОVВ)', '23.01.2015', '05:55', 'Москва (Шереметьево)(SVО) ', 'SU1549'
        в ->>> 'Новосибирск', 'Москва', '23.01.2015', '05:55'
        :param for_graph:  если True - то в городах где есть - добавим перед тире \n, что бы вместить в узел графа
        :return: список 'Новосибирск', 'Москва', '23.01.2015'"""
        lst = []
        for city_src, date_depart, time_depart, city_dst, fly in self.person_trips:
            # оставим от города только "города"     Ростов-на-Дону(РОВ) --> Ростов-на-Дону
            re_res = re.findall('(\w+\-?\w+\-?\w+)', city_src)
            src = re_res[0] if len(re_res) else city_src
            re_res = re.findall('(\w+\-?\w+\-?\w+)', city_dst)
            dst = re_res[0] if len(re_res) else city_dst
            # добавим в список, удалив всякие выражения вида -ЛАДОЖ
            lst.append(
                (self.replace_bad_words(src, for_graph), self.replace_bad_words(dst, for_graph),
                 date_depart, time_depart))
        return lst

    @staticmethod
    def get_days_differ(date1, date2):
        """Получим разницу в датах в днях"""
        date_format = '%d.%m.%Y'
        result = -1
        try:
            d1 = datetime.strptime(date1, date_format)
            d2 = datetime.strptime(date2, date_format)
        except Exception as e:
            print(datetime.now(), e)
        else:
            result = (d2 - d1).days
        return result

    def is_russia(self, city):
        """ Проверим по yandex API является ли этот город Российским
        :param city: город по которому проверяем
        :return: 1 - РФ, 0 - не РФ
        """
        # получим список уже проверенных городов
        cities = self.checked_cities.keys()
        # если проверяемый город есть в этом списке - сразу получим результат
        if city in cities:
            return self.checked_cities[city]
        else:
            url = f'https://geocode-maps.yandex.ru/1.x/?geocode={city}&format=json&key=ACNZFlwBAAAAut9LZQIAKQvqWvGQPsv8yOOAO1cd04yRlfEAAAAAAAAAAABL7YYFTBZ_yLbPTTS9mPggxujm2Q==&languege=ru'
            try:
                resp = requests.get(url)
                j = json.loads(resp.text)

                # обработаем json
                country = ''
                # получим список geo-объектов
                geo_object_list = j['response']['GeoObjectCollection']['featureMember']
                # #
                if geo_object_list:
                    country = geo_object_list[0]['GeoObject']['description']

                result = 1 if 'Россия' in country else 0
            except Exception as e:
                print(datetime.now(), e)
                result = -1

            # сохраним информацию о городе
            self.checked_cities[city] = result
            return result

    def check_abroad(self):
        """
        Проверим был ли человек где то долго за границей, проверять будет по следующей логике:
        проверяю город отправления - если он не РФ, проверяю город прибытия в следующей строке
        что бы они были одинаковыми и высчитываю дату пребывания
        :return: список с городами в которых был за границей и дата Минск, 20.11.2017, 25.11.2017"""
        # если в настройках указаны "заграничные дни"
        if self.days_aboard > 0:
            # начнем парсить полученный список
            for i, (src, dst, date_depart, time) in enumerate(self.trips_cities):
                # проверим город назначения, является ли он заграгницей и если есть еще строки
                if self.is_russia(dst) == 0 and i + 1 < len(self.trips_cities):
                    # получим день отправления следующего рейса
                    next_src = self.trips_cities[i + 1][0]
                    # получим дату следующего вылета для расчета дней пребывания
                    next_date = self.trips_cities[i + 1][2]
                    # сравним город отправления, он должен быть таким же как текущий отправления
                    # и сколько там был по дням
                    if dst == next_src and self.get_days_differ(date_depart, next_date) >= self.days_aboard:
                        self.abroad_trips.append((dst, date_depart, next_date))

    @staticmethod
    def prepare_sorted_trips(tripsall):
        """1. Уберем из списка дату и время
        2. Получим количество поездок по каждому направлению
        3. Отсортируем список
        4. удалим повторы
        :param tripsall: список поездок вида [(src, dst, date, time)...
        :return: список вида (Москва, Псков, N)"""
        # удалим из нашего списка дату и время
        tripsall = [(src, dst) for src, dst, _, _ in tripsall]
        # количество поездок по каждому направлению
        c_dict = collections.Counter(tripsall)
        # получим список вида [ (Москва, Псков, N), .. ], где N количество поездок по данному направлению
        trips = [(src, dst, c_dict[(src, dst)]) for src, dst in c_dict]
        # отсортируем по количеству поездок по направлению
        return list(sorted(trips, key=lambda x: x[2], reverse=True))

    def prepare_plots_data(self):
        """Подготовим информацию длч работы с графикой"""
        # получим список городов для графа, в нем после тире(-) будет перенос коретки
        trips = self.get_cities(for_graph=True)
        # Переведем все в большой регистр
        trips = [(str(src).upper(), str(dst).upper(), date, time) for src, dst, date, time in trips]
        # удалим повторы
        self.trips_all = list(set(trips))

        # 1. Получим грани с весом
        # отсортируем по количеству поездок по направлению
        self.trips_sorted = self.prepare_sorted_trips(self.trips_all)

        # 2. Получим список городов и количество посещений этих городов
        # 2.1. сформируем список городов С ПОВТОРАМИ
        cities_uniq = []
        for src, dst, _, _ in self.trips_all:
            cities_uniq.append(src)
            cities_uniq.append(dst)
        # произведем расчет повторений каждого города
        self.cities_dict = collections.Counter(cities_uniq)
        # отсортируем словарь с количеством, в итого получим список вида [(МСК, 81), (СПБ, 11)...]
        self.cities_sorted = sorted(self.cities_dict.items(), key=lambda kv: kv[1], reverse=True)

        # 3. подготовим список поездок совершенный в праздничные дни
        # 30.12 - 01.01, 22-24.02, 7-9.03, 8-12.05
        self.trips_holidays_all = []
        for src, dst, date, time in self.trips_all:
            # проверим входит ли дата в праздничные
            dt = datetime.strptime(date, '%d.%m.%Y')
            m = dt.month
            d = dt.day
            if (m == 12 and (30 <= d <= 31)) or (m == 1 and (1 <= d <= 10)) or (m == 2 and (22 <= d <= 24)) \
                    or (m == 3 and (7 <= d <= 9)) or (m == 5 and (8 <= d <= 12)):
                self.trips_holidays_all.append((src, dst, date, time))
        # 3.1. отсортируем праздничные поездки
        self.trips_holidays_all = sorted(self.trips_holidays_all, key=lambda x: datetime.strptime(x[2], '%d.%m.%Y'),
                                         reverse=True)
        # 3.2. сформируем список городов праздничных дней
        if len(self.trips_holidays_all):
            cities_uniq = []
            for src, dst, _, _ in self.trips_holidays_all:
                cities_uniq.append(src)
                cities_uniq.append(dst)
            # произведем расчет повторений каждого города
            cities_dict = collections.Counter(cities_uniq)
            # отсортируем словарь с количеством, в итого получим список вида [(МСК, 81), (СПБ, 11)...]
            self.cities_sorted_holidays = sorted(cities_dict.items(), key=lambda kv: kv[1], reverse=True)

            # 3.3. получим список поездок с количество каждого направления и отсортируем его
            self.trips_sorted_holidays = self.prepare_sorted_trips(self.trips_holidays_all)

        self.trips_weekends_all = []
        # 4. подготовим список поездок совершенный в выходные дни (пятница, суббота, воскр)
        for src, dst, date, time in self.trips_all:
            # проверим входит ли дата в праздничные
            dt = datetime.strptime(date, '%d.%m.%Y')
            if dt.weekday() in (5, 6, 7):
                self.trips_weekends_all.append((src, dst, date, time))

        self.trips_weekends_all = sorted(self.trips_weekends_all, key=lambda x: datetime.strptime(x[2], '%d.%m.%Y'),
                                         reverse=True)
        # 4.1. сформируем список городов выходных дней
        if len(self.trips_weekends_all):
            cities_uniq = []
            for src, dst, _, _ in self.trips_weekends_all:
                cities_uniq.append(src)
                cities_uniq.append(dst)
            # произведем расчет повторений каждого города
            cities_dict = collections.Counter(cities_uniq)
            # отсортируем словарь с количеством, в итого получим список вида [(МСК, 81), (СПБ, 11)...]
            self.cities_sorted_weekends = sorted(cities_dict.items(), key=lambda kv: kv[1], reverse=True)
            # 4.2 получим список поездок с количество каждого направления и отсортируем его
            self.trips_sorted_weekends = self.prepare_sorted_trips(self.trips_weekends_all)

    def draw_heatmap(self):
        """Построим heatmap c количеством поездок в дни недели"""
        # создадим фигуру из 4 частей
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 8))

        # 1. подготовим данные для heatmap
        # двумерный массив для хранения данных поездках в определенные дни
        trips_2d = np.zeros((7, 12), dtype='int')

        if len(self.person_trips):
            # начнем обрабатывать все поездки, пополняя наш массив
            for _, date, _, _, _ in self.person_trips:
                dt = datetime.strptime(date, '%d.%m.%Y')
                m = dt.month
                d = dt.weekday()
                trips_2d[d, m-1] += 1

            months = ('Январь', 'Февраль', 'Март', 'Апрель', 'Май',
                      'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь')

            days = ('Понедельник', 'Вторник', 'Среда', 'Четверг',
                    'Пятница', 'Суббота', 'Воскресенье')

            heatplot = ax1.imshow(trips_2d, cmap='Blues')
            # cbar = fig.colorbar(heatplot, ticks=[trips_2d.min(), trips_2d.max()])
            # ax1.plot(cbar)

            # We want to show all ticks...
            ax1.set_xticks(np.arange(len(months)))
            ax1.set_yticks(np.arange(len(days)))
            # ... and label them with the respective list entries
            ax1.set_xticklabels(months)
            ax1.set_yticklabels(days)

            # Rotate the tick labels and set their alignment.
            plt.setp(ax1.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")

            # Loop over data dimensions and create text annotations.
            for i in range(len(days)):
                for j in range(len(months)):
                    ax1.text(j, i, trips_2d[i, j], ha="center", va="center", color="w")

            ax1.set_title("Количество поездок в дни недели")
            fig.tight_layout()

            # 2. начнем рисовать пирог с поездками по дням недели

            # массив с суммой поездок в определенный день
            weekdays_sums = np.sum(trips_2d, axis=1)
            # количество всех поездок
            all_trips = np.sum(weekdays_sums)

            # если поездки вообще есть
            if all_trips > 0:
                # рассчитаем проценты
                percents = [round((cnt / all_trips * 100)) for cnt in weekdays_sums]
                day_names = [d for i, d in enumerate(days) if percents[i] > 0]
                # найдем максимальный элемент в поездках по дням
                max_trip = max(percents)
                percents = [p for p in percents if p > 0]
                explode = [0.1 if cnt == max_trip else 0 for cnt in percents]

                ax2.set_title("Распределение количества поездок по дням недели")
                # если есть значения меньше 2% то уберем надписи и добавим легенду - так будет наглядней
                small_cnt = [p for p in percents if p < 4]
                if len(small_cnt) > 2:
                    patches, texts, _ = ax2.pie(percents, explode=explode, autopct='%1.1f%%', startangle=90,
                                                textprops={'fontsize': 10})
                    ax2.legend(patches, day_names, loc="best")
                else:
                    ax2.pie(percents, explode=explode, labels=day_names, autopct='%1.1f%%', startangle=90,
                            textprops={'fontsize': 9})

                ax2.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

                # 3. построим bar c количеством поездок в каждый месяц
                # массив с суммой поездок в определенный месяц
                months_sums = np.sum(trips_2d, axis=0)
                percents = [round((cnt / all_trips * 100)) for cnt in months_sums]
                max_trip = max(percents)

                # подправим массивы для пирога так, что бы не выводить нулевые значения
                months = [m for i, m in enumerate(months) if percents[i] > 0]
                percents = [p for p in percents if p > 0]
                explode = [0.1 if p == max_trip else 0 for p in percents]

                ax4.set_title("Распределение количества поездок по месяцам")
                # если есть значения меньше 2% то уберем надписи и добавим легенду - так будет наглядней
                small_cnt = [p for p in percents if p < 4]
                if len(small_cnt) > 2:
                    patches, texts, _ = ax4.pie(percents, explode=explode, autopct='%1.1f%%', startangle=90,
                                                textprops={'fontsize': 10})
                    ax4.legend(patches, months, loc="best")
                else:
                    ax4.pie(percents, explode=explode, labels=months, autopct='%1.1f%%', startangle=90,
                            textprops={'fontsize': 9})

                ax4.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

            # 4. построим bar c количеством поездок в максимальный день
            # найдем максимальный день, в массиве найдем индекс где максимальный элемент
            ind = np.unravel_index(np.argmax(weekdays_sums, axis=None), weekdays_sums.shape)
            # в связи с тем, что ind это кортеж вида (столбец, строка), вытащим нулевой элемент картежа
            self.day_ind = ind[0]
            self.day = days[self.day_ind]
            # подправим окончание дня недели (Среда -> Среду)
            if self.day.endswith('а'):
                self.day = self.day.replace('a', 'у').lower()

            ax3.set_title(f'ТОП 5 поездок осуществленные в {self.day}')

            # получим список с поездками в этот день
            trips_in_topday = []
            for src, dst, date, time in self.trips_all:
                if datetime.strptime(date, '%d.%m.%Y').weekday() == self.day_ind:
                    trips_in_topday.append((src, dst))

            # подсчитаем поездки
            toptrips_dict = collections.Counter(trips_in_topday)
            # отсортируем словарь
            self.trips_sorted_topday = sorted(toptrips_dict.items(), key=lambda kv: kv[1], reverse=True)
            # оставим только 5 поездок
            toptrips = self.trips_sorted_topday[:5]

            # получим названия 5 наиболее частых поездок
            bars = []
            heights = []
            for src_dst, cnt in toptrips:
                src, dst = src_dst
                src = src.replace('\n', '')
                dst = dst.replace('\n', '')
                bars.append(f'{src.capitalize()}-\n{dst.capitalize()}')
                heights.append(cnt)

            y_pos = range(len(bars))
            colors = ['#d62728' if i == 0 else '#32628d' for i in range(len(bars))]
            rect = ax3.bar(y_pos, heights, color=colors)

            self.autolabel(rect, "center", ax3)
            fontdict = {'fontsize': 9}
            ax3.set_xticks(y_pos)
            ax3.set_xticklabels(bars, fontdict=fontdict)
            ax3.set_xlabel('', fontsize=8)

            plt.savefig('tmp/heatmap_pie_bars.png')
            plt.close(fig)

    @staticmethod
    def autolabel(rects, xpos, ax_):

        """
        Attach a text label above each bar in *rects*, displaying its height.

        *xpos* indicates which side to place the text w.r.t. the center of
        the bar. It can be one of the following {'center', 'right', 'left'}.
        """
        xpos = xpos.lower()  # normalize the case of the parameter
        ha = {'center': 'center', 'right': 'left', 'left': 'right'}
        offset = {'center': 0.5, 'right': 0.57, 'left': 0.43}  # x_txt = x + w*off
        fontdict = {'fontsize': 6}

        for rect in rects:
            height = rect.get_height()
            ax_.text(rect.get_x() + rect.get_width() * offset[xpos], 1.01 * height, '{}'.format(height),
                     ha=ha[xpos], va='bottom', fontdict=fontdict)

    def draw_bar_cities(self, ax, cities_sorted, cities_count, header):
        """Построим бар с городами, за исключением москвы
        :param ax: ax на котором рисуем бар
        :param cities_sorted: отсортировнных список с городами и их количеством [(Псков, 11)]
        :param cities_count: количество поездок которое отображаем
        :param header: заголовок
        """
        cnts = [cnt for name, cnt in cities_sorted]
        # оставим в списке города у которых процент >0
        cities = [(name, cnt) for name, cnt in cities_sorted[:10] if name != 'МОСКВА' and round((cnt / sum(cnts) * 100)) > 1]
        # получим нужные елементы из списка с городами
        cnts = [cnt for name, cnt in cities[:cities_count]]
        names = [name for name, cnt in cities[:cities_count]]

        # 2. строим бар
        y_pos = range(len(names))
        colors = ['#d62728' if i == 0 else '#32628d' for i in range(len(cnts))]
        rect1 = ax.bar(y_pos, cnts, color=colors)
        self.autolabel(rect1, "center", ax)
        fontdict = {'fontsize': 8}
        ax.set_title(f'{header}, ТОП {len(names)}*')
        ax.set_xticks(y_pos)
        ax.set_xticklabels(names, fontdict=fontdict)
        ax.set_xlabel('\n* - за исключением Москвы', fontsize=6)

    def draw_bar_trips(self, ax, trips_sorted, trips_count, header):
        """Построим бар с городами, за исключением москвы
        :param ax: ax на котором рисуем бар
        :param trips_sorted: отсортировнных список с городами и их количеством [(Псков, 11)]
        :param trips_count: количество поездок которое отображаем
        :param header: заголовок
        """
        heights = [cnt for _, _, cnt in trips_sorted[:trips_count]]
        bars = [f'{src.capitalize()}\n{dst.capitalize()}' for src, dst, _ in trips_sorted[:trips_count]]
        y_pos = range(len(bars))
        colors = ['#d62728' if i == 0 else '#32628d' for i in range(len(heights))]
        rect2 = ax.bar(y_pos, heights, color=colors)
        self.autolabel(rect2, "center", ax)
        ax.set_title(f'{header}, ТОП {len(bars)}')
        ax.set_xticks(y_pos)
        fontdict = {'fontsize': 8}
        ax.set_xticklabels(bars, fontdict=fontdict)
        # ax2.set_xticklabels(bars, fontdict=fontdict, rotation=8)

    def draw_bars_trips_cities(self):
        """1. Bar на основе самых частых поездках
        2. BAR c посещаемыми городами без москвы"""
        # количество поездок которое будем отображать на барах
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 4))
        # установим рассмотяние между
        plt.subplots_adjust(wspace=0.2)
        # 1, 2. Построим бар на основе посещения городов, 10 самых частых перелетов, исключив Москву,
        # 2. строим бар
        self.draw_bar_cities(ax1, self.cities_sorted, cities_count=5, header='Самые посещаемые города')
        # 3. Построим бар с направлениями
        self.draw_bar_trips(ax2, self.trips_sorted, trips_count=5, header='Самые частые поездки')
        # сохраняем фигуру
        # установим рассмотяние между
        plt.subplots_adjust(wspace=0.1)
        fig.savefig('tmp/pie_bars_cities.png')
        plt.close(fig)

    def draw_holidays_bars(self):
        """Построим графику для поездок совершенных в праздничные дни и выходные дни
        30.12 - 01.01, 22-24.02, 7-9.03, 8-12.05"""
        fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(15, 9))
        # fig.()
        # установим рассмотяние между
        plt.subplots_adjust(wspace=0.1, hspace=0.4)
        fig.suptitle('Поездки в праздничные и выходные дни', fontsize=12)

        # ax1 - поездки в праздничые дни
        # 1. количество посещений каждого top-города в праздничные дни
        self.draw_bar_cities(ax1, self.cities_sorted_holidays, 5, header='Города посещаемые в праздничные дни')
        # 1.2. Построим бар с top-направлениями в праздничные дни
        self.draw_bar_trips(ax2, self.trips_sorted_holidays, 5, header='Поездки в праздничные дни')
        # 2. количество посещений каждого top-города в выходные дни
        self.draw_bar_cities(ax3, self.cities_sorted_weekends, 5, header='Города посещаемые в выходные дни')
        # 2.1. Построим бар с top-направлениями в выходные дни
        self.draw_bar_trips(ax4, self.trips_sorted_weekends, 5, header='Поездки в выходные дни')

        # сохраняем фигуру
        fig.savefig('tmp/holiday_bars.png')
        plt.close(fig)

    def draw_graph(self):
        # Построим граф
        fig, ax = plt.subplots(figsize=(14, 5))
        # получим 10 самых популярных поездок
        edges_w = list(filter(lambda x: x[2] > 1, self.trips_sorted[:10]))
        # удалим веса граней
        edges = [(src, dst) for src, dst, _ in edges_w]

        if len(edges) > 0:
            fig.suptitle('Граф основных поездок', fontsize=12)
            # Для пмосквы сделаем фиксированный размер
            node_sizes_graph = []
            for key in self.cities_dict.keys():
                if key == 'Москва':
                    size = 5000
                else:
                    size = min(1700 + (self.cities_dict[key])*5, 5000)
                node_sizes_graph.append(size)

            graph = nx.MultiDiGraph()
            graph.add_edges_from(edges)
            # graph.add_weighted_edges_from(edges)
            # получим словарь в котором будет информация о количестве вхождейни в узел
            node_colors = ['#93e4fe' if graph.degree[node] < 4 else '#fa5a5b' for node in graph.nodes()]
            # подготовим label узлов
            # edge_labels = dict([((u, v,), d['weight']) for u, v, d in graph.edges(data=True)])
            edge_labels = dict([((u, v,), w) for u, v, w in edges_w])

            pos = nx.spring_layout(graph)
            nx.draw(graph, pos=pos, node_size=node_sizes_graph, with_labels=True, edge_color='grey', font_size=9,
                    # node_color=node_colors)
                    node_color=node_colors, edge_labels=edge_labels)
            # отрисуем значение весов
            nx.draw_networkx_edge_labels(graph, pos=pos, edge_labels=edge_labels, edges=edges, font_size=7)
            fig.savefig('tmp/graph.png')
        plt.close(fig)

    @staticmethod
    def parse_userinfo(buff):
        """получим информацию о пользователе, а именно имя, паспортные данные и дату рождения
        если парсер не сможет получить ничего кроме имени - выведет исходную сттроку"""
        result = buff
        name = re.findall('([a-zA-ZА-Яа-я]{4,}\s[a-zA-ZА-Яа-я]{4,}\s(?:[a-zA-ZА-Яа-я]{4,}\s)?)', buff)
        bd = re.findall('(?:\d{2}[\w]{3}\d{2}|\d{2}\.\d{2}\.\d{4})', buff)
        passports = re.findall('(?:\d{1,2}\s?\d{2}\s?№?\s\d{6}|\d?[A-ZА-Я]{2,5}\s?№?\s\d{6})', buff)
        if name and (bd or passports):
            result = name[0].lstrip(' ').rstrip(' ')
            if bd:
                result += f'\nДата рождения: {",".join(bd)}'
            if passports:
                result += f'\nПаспорт: {",".join(passports)}'
        return result

    @staticmethod
    def remove_sym(in_str, sym='\n'):
        """
        Удалим символ sym из str
        :param in_str: входная строка
        :param sym: удаляемый символ
        :return:
        """
        res = in_str
        if sym in in_str:
            lst = in_str.split(sym)
            res = ''.join(lst)
        return res

    def prepare_iad(self, document):
        """Подготовим аналитический документ, добавим его в документ переданыый по ссылке
        :param document: документ, в который впишем аналитику
        :return: измененный документ docx
        """
        if len(self.person_fio):
            self.person_fio = self.person_fio.replace('(', '\n(')
            document.add_heading(self.person_fio.rstrip('\n'), 0)

        # если есть графика
        # if os.path.exists('tmp/heatmap_pie_bars.png') or os.path.exists('tmp/graph_pie_bars.png'):
        if os.path.exists('tmp/heatmap_pie_bars.png'):
            document.add_picture('tmp/heatmap_pie_bars.png', height=Inches(5.2))

        # вставим поездки топ-дня
        if len(self.trips_sorted_topday) > 1:
            trips = self.trips_sorted_topday[:10]
            # оставим где больше 1
            if len(trips) > 1:
                trips_f = list(filter(lambda x: int(x[1]) > 1, trips))
                if len(trips_f) > 1:
                    trips = trips_f

            # вставим таблицу с поездками выходных дней
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font = paragraph.add_run(f'Основные поездки в {self.day}, ТОП {len(trips)}')
            font.bold = True
            # создадим таблицу
            table = document.add_table(rows=(len(trips) + 1), cols=2)
            table.allow_autofit = True
            # впишем заголовок
            row = table.rows[0]
            row.cells[0].width = Cm(1.5)
            row.cells[0].text = 'Количество\nпоездок'
            row.cells[1].text = 'Поездка\\Перелет'

            # пройдем все поездки
            for i, ((src, dst), cnt) in enumerate(trips):
                row = table.rows[i+1]
                row.cells[0].width = Cm(1.5)
                row.cells[0].text = str(cnt)
                row.cells[1].text = f'{self.remove_sym(src.capitalize())} - {self.remove_sym(dst.capitalize())}'
            document.add_page_break()

        if os.path.exists('tmp/pie_bars_cities.png'):
            document.add_picture('tmp/pie_bars_cities.png', height=Inches(2.6))
        if os.path.exists('tmp/graph.png'):
            document.add_picture('tmp/graph.png', height=Inches(3.7))

        # вставим таблицу с посещениями городов
        if len(self.cities_sorted) > 1:
            trips = self.cities_sorted[:10]
            if len(trips) > 1:
                trips_f = list(filter(lambda x: int(x[1]) > 1, trips))
                if len(trips_f) > 1:
                    trips = trips_f

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font = paragraph.add_run(f'Самые посещаемые города, ТОП {len(trips)}')
            font.bold = True

            # создадим таблицу
            table = document.add_table(rows=(len(trips) + 1), cols=2)
            table.allow_autofit = True
            # впишем заголовок
            row = table.rows[0]
            row.cells[0].width = Cm(1.5)
            row.cells[0].text = 'Количество\nпосещений'
            row.cells[1].text = 'Город'

            # пройдем все поездки
            for i, (city, cnt) in enumerate(trips):
                row = table.rows[i+1]
                row.cells[0].width = Cm(1.5)
                row.cells[0].text = str(cnt)
                row.cells[1].text = f'{self.remove_sym(city.capitalize())}'

        # вставим таблицу с частыми поездками
        # [(мск, пск, 12)...]
        if len(self.trips_sorted) > 1:
            trips = self.trips_sorted[:10]
            if len(trips) > 1:
                trips_f = list(filter(lambda x: int(x[2]) > 1, trips))
                if len(trips_f)>1:
                    trips = trips_f

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font = paragraph.add_run(f'Наиболее частые поездки, ТОП {len(trips)}')
            font.bold = True

            # создадим таблицу
            table = document.add_table(rows=(len(trips) + 1), cols=2)
            table.allow_autofit = True
            # впишем заголовок
            row = table.rows[0]
            row.cells[0].width = Cm(1.5)
            row.cells[0].text = 'Количество\nпоездок'
            row.cells[1].text = 'Маршрут'

            # пройдем все поездки
            for i, (src, dst, cnt) in enumerate(trips):
                row = table.rows[i + 1]
                row.cells[0].width = Cm(1.5)
                row.cells[0].text = str(cnt)
                row.cells[1].text = f'{self.remove_sym(src.capitalize())} - {self.remove_sym(dst.capitalize())}'

            document.add_page_break()

        # todo вставить таблицу с частыми поездками (топ 20)
        # todo в графе убрать веса и размеры сделать больше точ
        if os.path.exists('tmp/holiday_bars.png') and (len(self.trips_holidays_all) or len(self.trips_weekends_all)):
            document.add_picture('tmp/holiday_bars.png', height=Inches(5.9))
            if len(self.trips_holidays_all) or len(self.trips_holidays_all):
                document.add_page_break()
            # вставим таблицу с поездками праздничных дней

            if len(self.trips_holidays_all):
                # получим список с топ 10 поездок
                # [[src, dst, date, time]...
                trips = self.trips_holidays_all[:10]

                paragraph = document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                font = paragraph.add_run('Поездки в праздничные дни (30.12-01.01, 22-24.02, 7-9.03, 8-12.05)\n'
                                         f'(последние {len(trips)} поездок)')
                font.bold = True

                # создадим таблицу
                table = document.add_table(rows=(len(trips) + 1), cols=1)
                table.allow_autofit = True
                for i, (src, dst, date, time) in enumerate(trips):
                    row = table.rows[i]
                    trip = f'{i+1}) {self.remove_sym(src.capitalize())} - {self.remove_sym(dst.capitalize())},' \
                           f' {date} {time}'
                    row.cells[0].text = trip

            if len(self.trips_holidays_all):
                # получим список с топ 10 поездок
                # [[src, dst, date, time]...
                trips = self.trips_weekends_all[:10]
                # вставим таблицу с поездками выходных дней
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                font = paragraph.add_run(f'Поездки в пятницу и выходные дни (Пятница, Суббота, Воскресенье)\n'
                                         f'(последние {len(trips)} поездок)')
                font.bold = True
                # создадим таблицу
                table = document.add_table(rows=(len(trips) + 1), cols=1)
                table.allow_autofit = True

                for i, (src, dst, date, time) in enumerate(trips):
                    row = table.rows[i]
                    weekday = datetime.strptime(date, '%d.%m.%Y').weekday()
                    trip = f'{i+1}) {self.remove_sym(src.capitalize())} - {self.remove_sym(dst.capitalize())},' \
                           f' {date} {time} ({self.days_dict[weekday]})'
                    row.cells[0].text = trip

        # если есть заграничные города:
        if self.abroad_trips:
            # todo добавить определение страны, можно тоже по яндекс.api
            document.add_heading('Пребывание за границей (на основании информации о приобретенных билетах)', 0)
            table1 = document.add_table(rows=(len(self.abroad_trips) + 1), cols=2)
            table1.allow_autofit = True
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = 'Город'
            hdr_cells[1].text = 'Даты пребывания'
            for n, (city, date1, date2) in enumerate(self.abroad_trips):
                cells = table1.rows[n + 1].cells
                cells[0].text = city
                cells[1].text = f'{date1} - {date2}'
            document.add_page_break()
        # todo вставить все поездки для кучи
        return document

    def prepare_trevelers(self, document):
        """Подготовим информацию о попутчиках
        :param document: документ, в который впишем попутчиков
        :return: измененный документ docx"""
        if len(self.person_fio):
            document.add_heading(self.person_fio, 0)

        document.add_heading('Попутчики', 0)
        # рассчитаем количество строк в таблице
        rows_count = 0
        for key in self.fellows_trips.keys():
            # количество поездок попутчика
            rows_count += len(self.fellows_trips[key])
            # и строка с фамилией попутчика
            rows_count += 1
        # и строка с заголовком дата\поездки\города
        rows_count += 1

        table = document.add_table(rows=rows_count, cols=2)
        row_num = 0
        # пройдем наш словарь и заполним созданную таблицу
        for key in self.fellows_trips.keys():
            row = table.rows[row_num]
            row.cells[0].merge(row.cells[1])
            row_num += 1
            # запишем в первую строку ФИО попутчика
            cell_text = row.cells[0].paragraphs[0].add_run(self.parse_userinfo(key))
            cell_text.bold = True

            for date, trip in self.fellows_trips[key]:
                row = table.rows[row_num]
                row.cells[0].width = Cm(1.5)
                row.cells[0].text = date
                row.cells[1].text = trip
                row_num += 1

        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(22)

        return document

    def save_trevelers(self, document):
        """
        Сохраним документ с попутчиками
        :param document:
        :return:
        """
        fpath = f'{self.out_dir}/{pathlib.Path(fn).stem}_travelers.docx'
        document.save(fpath)
        if len(self.person_fio):
            print(f'{self.person_fio} has travelers:', fpath)
        else:
            print('He has travelers:', fpath)

    def save_iad(self, document):
        """Сохраним документ с аналитикой
        :param document:
        :return:"""
        fpath = f'{self.iad_dir}/{pathlib.Path(fn).stem}_iad.docx'
        document.save(fpath)

    def save_user_data(self):
        """сохраним полученные данные в word"""
        try:

            # сохраним в word
            document = Document()
            sections = document.sections
            section = sections[0]
            section.orientation = WD_ORIENTATION.LANDSCAPE

            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.header_distance = Inches(1)
            section.footer_distance = Inches(1)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # подготовим аналитику
            document = self.prepare_iad(document)
            # сохраним iad
            self.save_iad(document)
            if self.fellows_trips.keys():
                # подготовим информацию о попутчиках
                document = self.prepare_trevelers(document)
                # сохраним iad
                self.save_trevelers(document)

        except Exception as ex:
            print(datetime.now(), ex)

    def execute(self, filename):
        """Главная функция обработки файла
        :param filename: имя входного файла
        :return: результат обработки файла True\False
        """
        # установим текущее имя файла
        self.current_fn = filename

        # пробуем считать входной файл
        if self.load_data():
            # запустим функцию парсер
            self.find_fellows()
            # изменим список перелетов для работы с графикой и "заграницей"
            self.trips_cities = self.get_cities()
            # проверим заграничные рейсы
            self.check_abroad()
            # подготовим данные для рисования графики
            self.prepare_plots_data()
            # подготовим данные для рисования графики
            self.draw_bars_trips_cities()
            # heatmap
            self.draw_heatmap()
            self.draw_graph()
            self.draw_holidays_bars()
            # todo заливка данных в postgrees
            # сохраним полученные данные
            self.save_user_data()
        else:
            print(datetime.now(), 'Ошибка чтения файла ', filename)

    def load_settings(self):
        try:
            """загрузим настройки"""
            config = configparser.ConfigParser()
            config.read('settings.ini', encoding='windows-1251')

            self.in_dir = config.get('settings', "input_directory")
            # создадим входную директорию
            pathlib.Path(self.in_dir).mkdir(parents=True, exist_ok=True)

            self.out_dir = config.get('settings', "output_directory")
            # если выходная директория не указана, создадим ее внутри входной
            if len(self.out_dir) == 0:
                self.out_dir = os.path.join(self.in_dir, 'results_travelers')

            pathlib.Path(self.out_dir).mkdir(parents=True, exist_ok=True)

            self.iad_dir = config.get('settings', "iad_directory")
            # если выходная директория не указана, создадим ее внутри входной
            if len(self.iad_dir) == 0:
                self.iad_dir = os.path.join(self.in_dir, 'iad')
            pathlib.Path(self.iad_dir).mkdir(parents=True, exist_ok=True)

            self.trips_number = int(config.get('settings', 'trips_number'))
            self.days_aboard = int(config.get('settings', 'days_abroad'))

            pathlib.Path('tmp').mkdir(parents=True, exist_ok=True)
        except Exception as e:
            print(datetime.now(), e)
            return False
        else:
            return True

    def clear_tmp(self):
        paths = glob.glob('tmp/*.png')
        # paths = [os.path.join(r, fp) for r, _, fp in os.walk('tmp') if len(fp) > 0]
        for p in paths:
            try:
                os.remove(p)
            except:
                pass

    def delete_tmp(self):
        self.clear_tmp()
        try:
            os.rmdir('tmp')
        except:
            pass


if __name__ == '__main__':
    #  инициализируем клавный класс обработчика
    finder = FellowFinder()
    # загрузим настройки
    if finder.load_settings():
        # получим список файлов для обработки
        files = [os.path.join(root, fn) for root, sub, files in os.walk(finder.in_dir) if root.find('results') == -1 and root.find('iad') == -1 for fn in files if len(fn) > 0 and fn.find('~$') == -1 and fn[0] != '.']
        # начнем обрабатывать список файлов
        for fn in files:
            finder.clear_tmp()
            finder.execute(fn)
            finder.clear()
    else:
        print(datetime.now(), 'Ошибка загрузки файла настроек.')
        input()
    finder.delete_tmp()
